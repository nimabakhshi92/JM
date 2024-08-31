# from .resp import *
import copy
import datetime
from env import *
from .models import *
from .serializers import *
from rest_framework.views import APIView
from rest_framework import status, generics
from rest_framework_simplejwt.views import TokenObtainPairView, TokenRefreshView
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework import mixins, viewsets
from django.db.models import Max
from .pagination import *
from .views import *
from .permissions import *
from django.db.models import Count, Q, Prefetch, F
from django.db.models.functions import Greatest
from django.db import transaction
from .enums import *

from rest_framework.response import Response
from django.http import HttpResponse, HttpResponseNotFound, FileResponse
from docx import Document
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
from django.conf import settings


class BaseListCreateDestroyVS(viewsets.GenericViewSet, mixins.CreateModelMixin,
                              mixins.DestroyModelMixin, mixins.ListModelMixin):
    pass


class BaseCreateUpdateDestroyVS(viewsets.GenericViewSet, mixins.CreateModelMixin,
                                mixins.UpdateModelMixin, mixins.DestroyModelMixin):
    pass


class BaseListCreateUpdateDestroyVS(viewsets.GenericViewSet, mixins.ListModelMixin, mixins.CreateModelMixin,
                                    mixins.UpdateModelMixin, mixins.DestroyModelMixin):
    pass


class BaseListCreateRetrieveUpdateDestroyVS(viewsets.GenericViewSet,
                                            mixins.RetrieveModelMixin,
                                            mixins.ListModelMixin,
                                            mixins.CreateModelMixin,
                                            mixins.DestroyModelMixin,
                                            mixins.UpdateModelMixin):
    pass


class MyUserRegisterView(generics.CreateAPIView):
    permission_classes = []

    def create(self, request, *args, **kwargs):
        user_data = request.data
        user_data._mutable = True
        user_data['username'] = user_data.get('email')
        serialized = MyUserRegisterSerializer(data=user_data)

        data = {}

        if serialized.is_valid():
            user = serialized.save()

            data['id'] = user.id
            data['username'] = user.username
            data['email'] = user.email
            refresh = RefreshToken.for_user(user)
            data['token'] = {
                'token': str(refresh.access_token),
                'refresh_token': str(refresh),
                'expires_at': datetime.now() + refresh.access_token.lifetime

            }

            return Response(data, status=status.HTTP_201_CREATED)
        else:
            data = serialized.errors
        return Response(data, status=status.HTTP_400_BAD_REQUEST)


class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer


class MyTokenRefreshView(TokenRefreshView):
    serializer_class = MyTokenRefreshSerializer


class TableOfContentsView(APIView):
    permission_classes = [PublicContentPermission]

    def get(self, request):
        user_id = int(self.request.query_params.get('user_id', -1))
        request_user = self.request.user
        queryset = ContentSummaryTree.objects.all()

        if request_user.id == user_id:
            queryset = queryset.filter(narration__owner__id=user_id)
        elif user_id == -1:
            queryset = queryset.filter(
                models.Q(narration__owner__is_superuser=True) | models.Q(narration__owner=None))
        else:
            queryset = queryset.none()

        queryset = queryset.values(
            'alphabet', 'subject_1', 'subject_2', 'subject_3', 'subject_4', 'expression', 'summary'
        ).distinct().order_by('alphabet', 'subject_1', 'subject_2', 'subject_3', 'subject_4')

        data = {}
        for item in queryset:
            alphabet = item['alphabet']
            subject = item['subject_1']
            sub_subject = item['subject_2']
            subject_3 = item['subject_3']
            subject_4 = item['subject_4']
            expression = item['expression']
            summary = item['summary']
            if alphabet == 'ت' and subject == 'توحید':
                print(sub_subject)

            if alphabet not in data:
                data[alphabet] = {
                    'alphabet': alphabet,
                    'subjects': []
                }

            alphabet_data = data[alphabet]
            subject_data = next(
                (sub for sub in alphabet_data['subjects'] if sub['title'] == subject),
                None
            )

            if subject_data is None:
                subject_data = {
                    'title': subject,
                    'sub_subjects': []
                }
                alphabet_data['subjects'].append(subject_data)

            sub_subject_data = next(
                (sub for sub in subject_data['sub_subjects'] if sub['title'] == sub_subject),
                None
            )

            if sub_subject_data is None:
                sub_subject_data = {
                    'title': sub_subject,
                    'subjects_3': []
                }
                subject_data['sub_subjects'].append(sub_subject_data)

            subjects_3_data = next(
                (sub for sub in sub_subject_data['subjects_3'] if sub['title'] == subject_3),
                None
            )

            if subjects_3_data is None:
                subjects_3_data = {
                    'title': subject_3,
                    'subjects_4': []
                }
                sub_subject_data['subjects_3'].append(subjects_3_data)

            subjects_4_data = next(
                (sub for sub in subjects_3_data['subjects_4'] if sub['title'] == subject_4),
                None
            )

            if subjects_4_data is None:
                subjects_4_data = {
                    'title': subject_4,
                    'content': []
                }
                subjects_3_data['subjects_4'].append(subjects_4_data)

            content_data = {
                'expression': expression,
                'summary': summary
            }
            subjects_4_data['content'].append(content_data)

        serializer = AlphabetSerializer(data=list(data.values()), many=True)
        serializer.is_valid(raise_exception=True)

        return Response(serializer.validated_data)


class SurahTableOfContentsView(APIView):
    permission_classes = [PublicContentPermission]

    def get(self, request):
        user_id = int(self.request.query_params.get('user_id', -1))
        request_user = self.request.user
        queryset = NarrationSubjectVerse.objects.all()

        if request_user.id == user_id:
            queryset = queryset.filter(models.Q(content_summary_tree__narration__owner__id=user_id))
        elif user_id == -1:
            queryset = queryset.filter(
                models.Q(content_summary_tree__narration__owner__is_superuser=True) | models.Q(
                    content_summary_tree__narration__owner=None))
        else:
            queryset = queryset.none()

        queryset = queryset.filter(content_summary_tree__alphabet='بیان').values(
            'quran_verse__surah_no', 'quran_verse__surah_name', 'quran_verse__verse_no', 'quran_verse__verse_content',
            'content_summary_tree__subject_3', 'content_summary_tree__subject_4',
            'content_summary_tree__subject_2', 'content_summary_tree__expression', 'content_summary_tree__summary',
        ).distinct().order_by('quran_verse__surah_no', 'quran_verse__verse_no',
                              'content_summary_tree__subject_2', 'content_summary_tree__subject_3',
                              'content_summary_tree__subject_4')

        data = {}
        for item in queryset:
            surah_no = item['quran_verse__surah_no']
            surah_name = item['quran_verse__surah_name']
            verse_no = item['quran_verse__verse_no']
            verse_content = item['quran_verse__verse_content']
            sub_subject = item['content_summary_tree__subject_2']
            subject_3 = item['content_summary_tree__subject_3']
            subject_4 = item['content_summary_tree__subject_4']
            expression = item['content_summary_tree__expression']
            summary = item['content_summary_tree__summary']

            if surah_no not in data:
                data[surah_no] = {
                    'surah_no': surah_no,
                    'surah_name': surah_name,
                    'verses': []
                }

            surah_data = data[surah_no]
            verses_data = next(
                (sub for sub in surah_data['verses'] if sub['verse_no'] == verse_no),
                None
            )

            if verses_data is None:
                verses_data = {
                    'verse_no': verse_no,
                    'verse_content': verse_content,
                    'sub_subjects': []
                }
                surah_data['verses'].append(verses_data)

            sub_subject_data = next(
                (sub for sub in verses_data['sub_subjects'] if sub['title'] == sub_subject),
                None
            )

            if sub_subject_data is None:
                sub_subject_data = {
                    'title': sub_subject,
                    'subjects_3': []
                }
                verses_data['sub_subjects'].append(sub_subject_data)

            subjects_3_data = next(
                (sub for sub in sub_subject_data['subjects_3'] if sub['title'] == subject_3),
                None
            )

            if subjects_3_data is None:
                subjects_3_data = {
                    'title': subject_3,
                    'subjects_4': []
                }
                sub_subject_data['subjects_3'].append(subjects_3_data)

            subjects_4_data = next(
                (sub for sub in subjects_3_data['subjects_4'] if sub['title'] == subject_4),
                None
            )

            if subjects_4_data is None:
                subjects_4_data = {
                    'title': subject_4,
                    'content': []
                }
                subjects_3_data['subjects_4'].append(subjects_4_data)

            content_data = {
                'expression': expression,
                'summary': summary
            }
            subjects_4_data['content'].append(content_data)

        serializer = SurahSerializer(data=list(data.values()), many=True)
        serializer.is_valid(raise_exception=True)

        return Response(serializer.validated_data)


class VersesTableOfContentsView(APIView):
    permission_classes = [PublicContentPermission]

    def get(self, request):
        user_id = int(self.request.query_params.get('user_id', -1))
        request_user = self.request.user
        queryset = NarrationSubjectVerse.objects.all()

        if request_user.id == user_id:
            queryset = queryset.filter(models.Q(content_summary_tree__narration__owner__id=user_id))
        elif user_id == -1:
            queryset = queryset.filter(
                models.Q(content_summary_tree__narration__owner__is_superuser=True) | models.Q(
                    content_summary_tree__narration__owner=None))
        else:
            queryset = queryset.none()

        queryset = queryset.values(
            'content_summary_tree__alphabet', 'content_summary_tree__subject_1', 'content_summary_tree__subject_2',
            'content_summary_tree__subject_3', 'content_summary_tree__subject_4',
            'content_summary_tree__expression', 'content_summary_tree__summary'
        ).distinct().order_by('content_summary_tree__alphabet', 'content_summary_tree__subject_1',
                              'content_summary_tree__subject_2', 'content_summary_tree__subject_3',
                              'content_summary_tree__subject_4')

        data = {}
        for item in queryset:
            alphabet = item['content_summary_tree__alphabet']
            subject = item['content_summary_tree__subject_1']
            sub_subject = item['content_summary_tree__subject_2']
            subject_3 = item['content_summary_tree__subject_3']
            subject_4 = item['content_summary_tree__subject_4']
            expression = item['content_summary_tree__expression']
            summary = item['content_summary_tree__summary']

            if alphabet not in data:
                data[alphabet] = {
                    'alphabet': alphabet,
                    'subjects': []
                }

            alphabet_data = data[alphabet]
            subject_data = next(
                (sub for sub in alphabet_data['subjects'] if sub['title'] == subject),
                None
            )

            if subject_data is None:
                subject_data = {
                    'title': subject,
                    'sub_subjects': []
                }
                alphabet_data['subjects'].append(subject_data)

            sub_subject_data = next(
                (sub for sub in subject_data['sub_subjects'] if sub['title'] == sub_subject),
                None
            )

            if sub_subject_data is None:
                sub_subject_data = {
                    'title': sub_subject,
                    'subjects_3': []
                }
                subject_data['sub_subjects'].append(sub_subject_data)

            subjects_3_data = next(
                (sub for sub in sub_subject_data['subjects_3'] if sub['title'] == subject_3),
                None
            )

            if subjects_3_data is None:
                subjects_3_data = {
                    'title': subject_3,
                    'subjects_4': []
                }
                sub_subject_data['subjects_3'].append(subjects_3_data)

            subjects_4_data = next(
                (sub for sub in subjects_3_data['subjects_4'] if sub['title'] == subject_4),
                None
            )

            if subjects_4_data is None:
                subjects_4_data = {
                    'title': subject_4,
                    'content': []
                }
                subjects_3_data['subjects_4'].append(subjects_4_data)

            content_data = {
                'expression': expression,
                'summary': summary
            }
            subjects_4_data['content'].append(content_data)

        serializer = AlphabetSerializer(data=list(data.values()), many=True)
        serializer.is_valid(raise_exception=True)

        return Response(serializer.validated_data)


class NarrationVS(BaseListCreateRetrieveUpdateDestroyVS):
    permission_classes = [NarrationPermission]
    pagination_class = MyPagination

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return NarrationRetrieveSerializer
        elif self.action == 'update':
            return NarrationUpdateSerializer
        else:
            return NarrationSerializer

    def get_queryset(self):
        alphabet = self.request.query_params.get('alphabet', None)
        subject = self.request.query_params.get('subject', None)
        sub_subject = self.request.query_params.get('sub_subject', None)
        subject_3 = self.request.query_params.get('subject_3', None)
        subject_4 = self.request.query_params.get('subject_4', None)
        imam_name = self.request.query_params.get('imam_name', None)
        narration_name = self.request.query_params.get('narration_name', None)
        surah_name = self.request.query_params.get('surah_name', None)
        verse_no = self.request.query_params.get('verse_no', None)
        subjects_search = self.request.query_params.get('subjects_search', None)
        texts_search = self.request.query_params.get('texts_search', None)
        user_id = int(self.request.query_params.get('user_id', -1))
        request_user = self.request.user
        # queryset = Narration.objects.all().annotate(
        #     bookmarks_count=Count('bookmarks', filter=Q(bookmarks__user=request_user)))

        queryset = Narration.objects.all()

        if self.action == 'retrieve' or self.action == 'list':
            if request_user.id == user_id:
                queryset = queryset.filter(models.Q(owner__id=user_id))
            elif user_id == -1:
                queryset = queryset.filter(models.Q(owner__is_superuser=True) | models.Q(owner=None))
            else:
                queryset = queryset.none()
        if alphabet:
            queryset = queryset.filter(content_summary_tree__alphabet=alphabet).distinct()
        if subject:
            queryset = queryset.filter(content_summary_tree__subject_1=subject).distinct()
        if sub_subject:
            queryset = queryset.filter(content_summary_tree__subject_2=sub_subject).distinct()
        if subject_3:
            queryset = queryset.filter(content_summary_tree__subject_3=subject_3).distinct()
        if subject_4:
            queryset = queryset.filter(content_summary_tree__subject_4=subject_4).distinct()
        if imam_name:
            queryset = queryset.filter(imam__name=imam_name)
        if narration_name:
            queryset = queryset.filter(name=narration_name)
        if surah_name:
            queryset = queryset.filter(content_summary_tree__verse__quran_verse__surah_name=surah_name).distinct()
        if verse_no:
            queryset = queryset.filter(content_summary_tree__verse__quran_verse__verse_no=verse_no).distinct()
        if subjects_search:
            for or_subject_list in subjects_search.split(','):
                queryset = queryset.filter(subjects__subject__in=or_subject_list.split('|'))
        if texts_search:
            original_queryset = queryset
            for or_text_list_str in texts_search.split(','):
                filtered_queryset = []
                or_text_list = or_text_list_str.split('|')
                for item in queryset:
                    for text in or_text_list:
                        if remove_arabic_characters(text) in remove_arabic_characters(item.content):
                            filtered_queryset.append(item)
                            break
                queryset = filtered_queryset
            queryset = original_queryset.filter(id__in=map(lambda x: x.id, queryset))

        queryset = queryset.prefetch_related(
            'subjects', 'footnotes', 'narration_verses', 'content_summary_tree').prefetch_related(
            Prefetch("bookmarks", queryset=Bookmark.objects.filter(user=request_user))
        ).annotate(
            bookmarks_count=Count('bookmarks', filter=Q(bookmarks__user=request_user)),
            content_summary_tree_last_modified=Max('content_summary_tree__modified'),
            # subjects_last_modified=Max('subjects__modified'),
            footnotes_last_modified=Max('footnotes__modified'),
            # narration_verses_last_modified=Max('narration_verses__modified'),
            last_modified=Greatest(
                'modified',
                'content_summary_tree_last_modified',
                # 'subjects_last_modified',
                'footnotes_last_modified',
                # 'narration_verses_last_modified',
            )
        )

        sort_by = self.request.query_params.get('sort_by', 'modified')
        if sort_by == 'modified':
            sort_by = 'last_modified'
        sort_type = self.request.query_params.get('sort_type', None)
        sort_type = '' if sort_type == 'asc' else '-'

        return queryset.distinct().order_by(f'{sort_type}{sort_by}')

    def create(self, request, *args, **kwargs):
        user = request.user.id
        data = request.data
        data['user_id'] = user
        serialized = NarrationSerializer(data=data)
        if serialized.is_valid():
            serialized.save()
            return Response(serialized.data, status=status.HTTP_201_CREATED)


class BookVS(mixins.ListModelMixin, mixins.CreateModelMixin, mixins.RetrieveModelMixin, mixins.UpdateModelMixin,
             viewsets.GenericViewSet):
    permission_classes = [PublicContentPermission]
    serializer_class = BookSerializer
    queryset = Book.objects.all()


class ImamVS(mixins.ListModelMixin, viewsets.GenericViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = ImamSerializer
    queryset = Imam.objects.all()


class SubjectVS(BaseListCreateUpdateDestroyVS):
    permission_classes = [PublicContentPermission]
    serializer_class = NarrationSubjectModelSerializer
    queryset = NarrationSubject.objects.all()

    def list(self, request, *args, **kwargs):
        subjects = NarrationSubject.objects.values_list('subject', flat=True).distinct()
        serializer = NarrationSubjectListSerializer({'subjects': subjects})
        return Response(serializer.data)


class QuranSurahVS(viewsets.GenericViewSet, mixins.ListModelMixin):
    permission_classes = [IsAuthenticated]
    serializer_class = QuranSurahSerializer
    queryset = QuranVerse.objects.values('surah_no', 'surah_name').annotate(no_of_verses=Max('verse_no'))


class QuranVS(viewsets.GenericViewSet, mixins.ListModelMixin):
    permission_classes = [PublicContentPermission]
    serializer_class = QuranVerseSerializer

    def get_queryset(self):
        surah_no = self.request.query_params.get('surah_no')
        verse_no = self.request.query_params.get('verse_no')

        queryset = QuranVerse.objects.all()
        if surah_no:
            queryset = queryset.filter(surah_no=surah_no)
        if verse_no:
            queryset = queryset.filter(verse_no=verse_no)

        return queryset


class NarrationFootnoteVS(BaseCreateUpdateDestroyVS):
    queryset = NarrationFootnote.objects.all()
    serializer_class = NarrationFootnoteSerializer
    permission_classes = [NarrationRelatedFieldPermission]


class ContentSummaryTreeVS(BaseListCreateUpdateDestroyVS):
    serializer_class = ContentSummaryTreeSerializer
    queryset = ContentSummaryTree.objects.all()
    permission_classes = [NarrationRelatedFieldPermission]


class FilterOptionsVS(generics.ListAPIView):
    serializer_class = FilterOptionsSerializer
    permission_classes = [PublicContentPermission]
    queryset = Narration.objects.all().values(
        # 'name',
        #                                       'imam__name',
        'content_summary_tree__alphabet',
        'content_summary_tree__subject_1',
        'content_summary_tree__subject_2',
        # 'content_summary_tree__subject_3',
        # 'content_summary_tree__subject_4',
        # 'content_summary_tree__verse__quran_verse__surah_name',
        # 'content_summary_tree__verse__quran_verse__verse_no',
        # 'content_summary_tree__verse__quran_verse__verse_content'
    ).distinct()


def word_is_in_splited_text(word, splited):
    if word == ' ':
        return -1
    elif word in splited:
        return 1
    else:
        return 0


from fuzzywuzzy import fuzz
from .views import remove_arabic_characters


def find_all(a_str, sub):
    start = 0
    while True:
        start = a_str.find(sub, start)
        if start == -1: return
        yield start
        start += len(sub)


def extract_arabic_text(content, lang_splitter):
    len_splitter = len(lang_splitter)
    indices = list(find_all(content, lang_splitter))
    arabic_parts = [content[(indices[i] + len_splitter):indices[i + 1]] for i in range(len(indices)) if
                    i % 4 == 0 and i < len(indices) - 1]
    arabic_text = ''.join(arabic_parts)
    return arabic_text


def is_similar(expression, content, tolerance):
    similarity_score = fuzz.partial_ratio(expression.lower(), content.lower())
    return similarity_score >= tolerance * 100


class SimilarNarrations(APIView):
    permission_classes = [PublicContentPermission]
    lang_splitter = 'ظظظ'
    tolerance = 0.7

    def post(self, request):

        text = request.data.get('text')
        if not text:
            return Response(data={'message': 'Narration text must be nonempty'}, status=400)

        request_user = request.user
        text_words = list(filter(lambda x: len(x) > 1, text.split(' ')))
        queryset = Narration.objects.filter(Q(owner=request_user) | Q(owner__is_superuser=True))

        original_queryset = queryset
        text = remove_arabic_characters(text)
        similar_narrations = []
        for narration in queryset:
            try:
                narration_arabic_text = extract_arabic_text(narration.content, self.lang_splitter)
                narration_arabic_text = remove_arabic_characters(narration_arabic_text)
                if is_similar(text, narration_arabic_text, self.tolerance):
                    similar_narrations.append(narration)
            except:
                continue

            # splited_narration = narration.content.split(' ')
            # intersection = [word_is_in_splited_text(word, splited_narration) for word in text_words]
            # intersection = np.array(intersection)
            # intersection_percent = sum(intersection == 1) / sum(intersection != -1) * 100
            #
            # reverse_intersection = [word_is_in_splited_text(word, text_words) for word in splited_narration]
            # reverse_intersection = np.array(reverse_intersection)
            # reverse_intersection_percent = sum(reverse_intersection == 1) / sum(reverse_intersection != -1) * 100

            # if intersection_percent > 70 or reverse_intersection_percent > 35:
            #     similar_narrations.append(narration)
        queryset = original_queryset.filter(id__in=map(lambda x: x.id, similar_narrations))
        queryset.distinct().order_by('-modified')

        serializer = NarrationSerializer(queryset, many=True)
        return Response(serializer.data)


class BookmarkVS(BaseListCreateDestroyVS):
    permission_classes = [IsAuthenticated]
    serializer_class = BookmarkSerializer

    def get_queryset(self):
        queryset = Bookmark.objects.filter(user=self.request.user)
        return queryset

    def create(self, request, *args, **kwargs):
        user = request.user.id
        data = request.data
        data['user_id'] = user
        serialized = BookmarkSerializer(data=data)
        if serialized.is_valid():
            serialized.save()
            return Response(serialized.data, status=status.HTTP_201_CREATED)


def is_valid_shared_narration_update(current_status, new_status):
    if current_status == SharedNarrationsStatus.PENDING.value:
        return new_status in [SharedNarrationsStatus.CHECKING.value]

    if current_status == SharedNarrationsStatus.CHECKING.value:
        return new_status in [SharedNarrationsStatus.REJECTED.value, SharedNarrationsStatus.ACCEPTED.value]

    if current_status == SharedNarrationsStatus.REJECTED.value:
        return new_status in [SharedNarrationsStatus.PENDING.value]

    if current_status == SharedNarrationsStatus.ACCEPTED.value:
        return new_status in [SharedNarrationsStatus.REJECTED.value,
                              SharedNarrationsStatus.TRANSFERRED.value]
    return False


class SharedNarrationsVS(BaseListCreateRetrieveUpdateDestroyVS):
    permission_classes = [SharedNarrationsPermission]
    serializer_class = SharedNarrationsSerializer

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['request'] = self.request
        return context

    def get_queryset(self):
        request_user = self.request.user

        # return SharedNarrations.objects.all()
        if is_checker_admin(request_user):
            return SharedNarrations.objects.all()
        if is_a_non_checker_admin(request_user):
            return SharedNarrations.objects.filter(sender=request_user)

    def get_serializer_class(self):
        if self.action == 'partial_update':
            return SharedNarrationsPatchSerializer
        return SharedNarrationsSerializer

    def create(self, request, *args, **kwargs):
        sender_id = request.user.id
        request_data = request.data
        data = copy.deepcopy(request_data)
        data['sender_id'] = sender_id

        receiver_id = checker_admin_id

        data['receiver_id'] = receiver_id
        serialized = SharedNarrationsSerializer(data=data)
        if serialized.is_valid():
            serialized.save()
            return Response(serialized.data, status=status.HTTP_201_CREATED)

    def partial_update(self, request, *args, **kwargs):
        request_user = request.user
        request_data = request.data
        data = copy.deepcopy(request_data)

        instance = SharedNarrations.objects.get(pk=kwargs.get('pk'))
        current_status = instance.status
        new_status = data.get('status')
        is_valid_update = is_valid_shared_narration_update(current_status, new_status)
        if not is_valid_update:
            return Response(status=status.HTTP_403_FORBIDDEN)

        if new_status == SharedNarrationsStatus.CHECKING.value:
            original_narration = instance.narration
            with transaction.atomic():
                new_narration = duplicate_narration(original_narration, request_user)
                data['receiver_narration'] = new_narration.id
                serialized = SharedNarrationsPatchSerializer(instance, data=data)
                if serialized.is_valid():
                    serialized.save()
                    return Response(serialized.data, status=status.HTTP_200_OK)
                else:
                    raise Exception('Bad request')

        if new_status == SharedNarrationsStatus.REJECTED.value:
            with transaction.atomic():
                received_narration_id = instance.receiver_narration.id
                received_narration = Narration.objects.get(pk=received_narration_id)
                received_narration.delete()
                return super().partial_update(request, *args, **kwargs)

        if new_status == SharedNarrationsStatus.ACCEPTED.value or new_status == SharedNarrationsStatus.PENDING.value:
            return super().partial_update(request, *args, **kwargs)


def duplicate_narration(original_narration, new_owner):
    new_narration = Narration.objects.create(
        name=original_narration.name,
        imam=original_narration.imam,
        narrator=original_narration.narrator,
        content=original_narration.content,
        book=original_narration.book,
        book_vol_no=original_narration.book_vol_no,
        book_page_no=original_narration.book_page_no,
        book_narration_no=original_narration.book_narration_no,
        owner=new_owner,
    )
    for subject in original_narration.subjects.all():
        NarrationSubject.objects.create(narration=new_narration, subject=subject.subject)

    for footnote in original_narration.footnotes.all():
        NarrationFootnote.objects.create(narration=new_narration,
                                         expression=footnote.expression,
                                         explanation=footnote.explanation
                                         )

    for cst in original_narration.content_summary_tree.all():
        new_cst = ContentSummaryTree.objects.create(narration=new_narration,
                                                    alphabet=cst.alphabet,
                                                    subject_1=cst.subject_1,
                                                    subject_2=cst.subject_2,
                                                    subject_3=cst.subject_3,
                                                    subject_4=cst.subject_4,
                                                    expression=cst.expression,
                                                    summary=cst.summary,
                                                    )
        try:
            quran_verse = cst.verse.quran_verse
            NarrationSubjectVerse.objects.create(quran_verse=quran_verse,
                                                 content_summary_tree=new_cst)
        except:
            pass

    for verse in original_narration.narration_verses.all():
        NarrationVerse.objects.create(narration=new_narration, quran_verse=verse.quran_verse)

    return new_narration


class DuplicateNarrationVS(APIView):
    permission_classes = [DuplicateNarrationPermission]

    def post(self, request, *args, **kwargs):
        request_user = request.user
        narration_id = kwargs.get('narration_id')
        original_narration = Narration.objects.get(pk=narration_id)

        with transaction.atomic():
            new_narration = duplicate_narration(original_narration, request_user)

            return Response(NarrationSerializer(new_narration).data, status=status.HTTP_201_CREATED)


def move_narration_to_main_site(*args, **kwargs):
    new_owner = User.objects.get(pk=super_admin_id)
    narration_id = kwargs.get('narration_id')
    original_narration = Narration.objects.get(pk=narration_id)

    with transaction.atomic():
        new_narration = duplicate_narration(original_narration, new_owner=new_owner)
        original_narration.delete()
        return new_narration


class MoveToMainSiteNarrationVS(APIView):
    permission_classes = [DuplicateNarrationPermission]

    def post(self, request, *args, **kwargs):
        with transaction.atomic():
            narration_id = kwargs.get('narration_id')
            original_narration = Narration.objects.get(pk=narration_id)

            if hasattr(original_narration, 'shared_narration'):
                shared_narration = original_narration.shared_narration
                shared_narration.status = SharedNarrationsStatus.TRANSFERRED.value
                shared_narration.save()
            new_narration = move_narration_to_main_site(*args, **kwargs)

            return Response(NarrationSerializer(new_narration).data, status=status.HTTP_201_CREATED)


class DownloadNarrationVS(APIView):
    content_color = RGBColor(16, 40, 190)
    translation_color = RGBColor(0, 0, 0)
    narrator_color = RGBColor(165, 45, 45)
    verse_color = RGBColor(14, 130, 8)
    name_color = RGBColor(255, 0, 0)
    address_color = RGBColor(160, 160, 160)

    general_font = "B Mitra"
    main_text_font = general_font
    translation_font = general_font
    expression_font = general_font
    name_font = general_font
    address_font = general_font
    narrator_font = general_font

    permission_classes = [PublicContentPermission]

    def set_run_format(self, run, font_name='Arial', font_size=12, bold=False, italic=False, color=None,
                       highlight=None):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if color:
            run.font.color.rgb = color
        if highlight:
            run.font.highlight_color = highlight

    def create_narration_docx(self, story):
        doc = Document()

        # Add title
        title_run = doc.add_paragraph()
        title_run.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        title_run = title_run.add_run(story.name)
        self.set_run_format(title_run, font_name=self.name_font, font_size=16, bold=True, color=self.name_color)

        vol_no_persian_name = 'جلد'
        book_vol_no = f'{vol_no_persian_name} {story.book_vol_no}' if story.book_vol_no else " "

        page_persian_name = 'صفحه'
        book_page = f'{page_persian_name} {story.book_page_no}' if story.book_page_no else " "
        narration_address = story.book.name + ' ' + book_vol_no + ' ' + book_page
        address_run = doc.add_paragraph()
        address_run.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        address_run = address_run.add_run(narration_address)
        # address_run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # address_run.add_run(narration_address)
        self.set_run_format(address_run, font_name=self.address_font, font_size=10, bold=False,
                            color=self.address_color)

        # Add narrator
        narrator_run = doc.add_paragraph()
        narrator_run.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        narrator_run = narrator_run.add_run(story.narrator.replace('#', ''))
        self.set_run_format(narrator_run, font_name=self.narrator_font, font_size=10, bold=False,
                            color=self.narrator_color)

        # Add content
        # content_run = doc.add_paragraph().add_run(story.content)
        # self.set_run_format(content_run, font_size=12, color=RGBColor(0, 0, 0))
        self.add_narration_content(doc, story.content)

        # Add footnotes
        # footnote_run = doc.add_paragraph().add_run(f"Footnotes: {story.footnotes}")
        # self.set_run_format(footnote_run, font_size=10, italic=True, color=RGBColor(128, 128, 128))

        # Add tags
        # tags_run = doc.add_paragraph().add_run(f"Tags: {', '.join(story.tags)}")
        # self.set_run_format(tags_run, font_size=10, italic=True, color="FF0000")

        # Add page numbers
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        page_num_run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_num_run._r.append(fldChar)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io

    def add_narration_content(self, doc, content):
        # Define colors
        content = content.replace('(', 'EEEEEE')
        content = content.replace(')', '333333')

        content = content.replace('«', 'DDDDDD')
        content = content.replace('»', 'JJJJJJ')

        content = content.replace('333333', '(')
        content = content.replace('EEEEEE', ')')

        content = content.replace('DDDDDD', '»')
        content = content.replace('JJJJJJ', '«')

        blue_color = self.content_color
        black_color = self.translation_color
        green_color = self.verse_color

        main_text_font = self.main_text_font
        translation_font = self.translation_font
        expression_font = self.expression_font

        # Step 1: Split the content based on "aaa"
        segments = content.replace('@', ' ').split("ظظظ")
        is_main_text = True  # Start with the assumption that the first segment is the main text
        rlm = chr(0x200F)
        # Step 2: Process each segment and format it accordingly
        for segment in segments:
            if not segment.strip():
                continue

            if is_main_text:
                # Create a new paragraph for the main text
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # Create a new paragraph for the translation
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Handle text wrapped in $
            parts = re.split(r'(\$[^$]+\$)', segment)
            for part in parts:
                if not part.strip():
                    continue

                # Adjust symbols within the text for RTL
                # part = re.sub(r'(\()', rf'{rlm}\1', part)  # Add RLM before (
                # part = re.sub(r'(\))', rf'\1{rlm}', part)  # Add RLM after )
                # part = re.sub(r'(<<)', rf'{rlm}\1', part)  # Add RLM before <<
                # part = re.sub(r'(>>)', rf'\1{rlm}', part)  # Add RLM after >>

                if part.startswith('$') and part.endswith('$'):
                    # Text wrapped in $
                    run = paragraph.add_run(part.strip('$'))
                    run.font.color.rgb = green_color
                    run.font.name = expression_font
                    run.font.size = Pt(14)
                else:
                    # Main text or translation
                    run = paragraph.add_run(part)
                    run.font.size = Pt(14) if is_main_text else Pt(12)
                    run.font.color.rgb = blue_color if is_main_text else black_color
                    run.font.name = main_text_font if is_main_text else translation_font

            is_main_text = not is_main_text  # Alternate for the next segment

    def get0(self, request):
        narration_ids_str = request.query_params.get('ids', None)
        if len(narration_ids_str):
            narration_ids_list = narration_ids_str.split(',')
            narrations = Narration.objects.filter(pk__in=narration_ids_list)
        else:
            narrations = Narration.objects.all()

        zip_buffer = BytesIO()
        try:
            with ZipFile(zip_buffer, 'w', ZIP_DEFLATED) as zip_file:
                parent = 'All - ' + datetime.now().strftime('%Y-%m-%d-%H-%M-%S')
                subjective_narrations = 'احادیث موضوعی'
                verses_interpretations = 'تفاسیر آیات'
                subjective_verses = 'آیات موضوعی'

                cst_unique_list = ContentSummaryTree.objects.filter(
                    narration__in=narrations
                ).values(
                    'alphabet', 'subject_1', 'subject_2'
                ).distinct().order_by('alphabet', 'subject_1', 'subject_2')
                for cst in cst_unique_list:
                    alphabet = cst['alphabet']
                    subject_1 = cst['subject_1']
                    subject_2 = cst['subject_2']

                    uncategorized = 'دسته بندی نشده'
                    category = alphabet or uncategorized
                    subcategory = subject_1 or uncategorized
                    subsubcategory = subject_2 or uncategorized

                    cst_narrations_id = narrations.filter(
                        content_summary_tree__alphabet=cst['alphabet'],
                        content_summary_tree__subject_1=cst['subject_1'],
                        content_summary_tree__subject_2=cst['subject_2']
                    ).values('id').distinct()

                    for index, narration_id in enumerate(cst_narrations_id):
                        narration = Narration.objects.get(pk=narration_id['id'])
                        doc_io = self.create_narration_docx(narration)

                        folder_path = os.path.join(parent, category, subcategory, subsubcategory)
                        filename = f"{str(index + 10001)[1:]}- {narration.name}.docx"
                        file_path_in_zip = os.path.join(folder_path, filename)
                        zip_file.writestr(file_path_in_zip, doc_io.getvalue())

                        narration_cst = narration.content_summary_tree

            # Set the buffer position to the beginning
            zip_buffer.seek(0)

            # Prepare the HttpResponse with the ZIP file content
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename=stories.zip'

            return response
        except:
            return HttpResponse(status=500)
        finally:
            # Close the buffer explicitly if necessary
            zip_buffer.close()

    def get(self, request):
        narration_ids_str = request.query_params.get('ids', None)
        if len(narration_ids_str):
            narration_ids_list = narration_ids_str.split(',')
            narrations = Narration.objects.filter(pk__in=narration_ids_list)
        else:
            narrations = Narration.objects.all()

        zip_buffer = BytesIO()
        try:
            with ZipFile(zip_buffer, 'w', ZIP_DEFLATED) as zip_file:
                parent = 'All - ' + datetime.now().strftime('%Y-%m-%d-%H-%M-%S')
                subjective_narrations = 'احادیث موضوعی'
                verses_interpretations = 'تفاسیر آیات'
                subjective_verses = 'آیات موضوعی'
                uncategorized = 'دسته بندی نشده'
                bayan = 'بیان'

                for index, narration in enumerate(narrations):
                    try:
                        doc_io = self.create_narration_docx(narration)
                        all_cst = narration.content_summary_tree.all()
                        if not len(all_cst):
                            folder_path = os.path.join(parent, uncategorized)
                            filename = f"{str(index + 10001)[1:]}- {narration.name}.docx"
                            file_path_in_zip = os.path.join(folder_path, filename)
                            zip_file.writestr(file_path_in_zip, doc_io.getvalue())
                            continue

                        processed_csts = []
                        verse_processed_csts = []
                        bayan_processed_csts = []
                        for cst in all_cst:
                            try:
                                alphabet = cst.alphabet
                                subject_1 = cst.subject_1
                                subject_2 = cst.subject_2

                                category = alphabet or uncategorized
                                subcategory = subject_1 or uncategorized
                                subsubcategory = subject_2 or uncategorized

                                category = category[:50]
                                subcategory = subcategory[:50]
                                subsubcategory = subsubcategory[:50]
                                filename = f"{str(index + 100001)[1:]}- {narration.name[:50]}.docx"

                                if not ([alphabet, subject_1, subject_2] in processed_csts) and not (alphabet == bayan):
                                    folder_path = os.path.join(parent, subjective_narrations, category, subcategory,
                                                               subsubcategory)
                                    file_path_in_zip = os.path.join(folder_path, filename)
                                    zip_file.writestr(file_path_in_zip, doc_io.getvalue())
                                    processed_csts.append([alphabet, subject_1, subject_2])

                                has_verse = hasattr(cst, 'verse')

                                if not has_verse:
                                    continue

                                if not ([alphabet, subject_1, subject_2] in verse_processed_csts) and not (
                                        alphabet == bayan):
                                    folder_path = os.path.join(parent, subjective_verses, category, subcategory,
                                                               subsubcategory)
                                    file_path_in_zip = os.path.join(folder_path, filename)
                                    zip_file.writestr(file_path_in_zip, doc_io.getvalue())
                                    verse_processed_csts.append([alphabet, subject_1, subject_2])

                                if not alphabet == bayan:
                                    continue

                                quran_verse = cst.verse.quran_verse

                                surah_name = quran_verse.surah_name
                                surah_no = quran_verse.surah_no
                                surah_desc = f'{surah_no}- {surah_name}'

                                verse_no = quran_verse.verse_no
                                verse_content = quran_verse.verse_content
                                verse_desc = f'{verse_no}- {verse_content[:40]}'

                                if not ([alphabet, surah_name, verse_no] in bayan_processed_csts):
                                    folder_path = os.path.join(parent, verses_interpretations, surah_desc,
                                                               verse_desc)
                                    file_path_in_zip = os.path.join(folder_path, filename)
                                    zip_file.writestr(file_path_in_zip, doc_io.getvalue())
                                    bayan_processed_csts.append([alphabet, surah_name, verse_no])
                            except:
                                pass
                    except:
                        pass

            # Set the buffer position to the beginning
            zip_buffer.seek(0)

            # Prepare the HttpResponse with the ZIP file content
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename=stories.zip'

            return response
        except:
            return HttpResponse(status=500)
        finally:
            # Close the buffer explicitly if necessary
            zip_buffer.close()


class DownloadNarrationBackupVS(viewsets.GenericViewSet, mixins.RetrieveModelMixin, mixins.ListModelMixin, ):
    def list(self, request, *args, **kwargs):
        folder_path = os.path.join( 'zipBackup')
        return Response([os.listdir(folder_path)[-1]], status=status.HTTP_200_OK)

    def retrieve(self, request, *args, **kwargs):
        folder_path = os.path.join( 'zipBackup')
        all_files = os.listdir(folder_path)
        all_files.sort()
        filename = all_files[-1]
        file_path = os.path.join('zipBackup', filename)
        if os.path.exists(file_path):
            return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
        else:
            return HttpResponseNotFound('File not found')

# class SpeedTestVS(APIView):
#     permission_classes = [IsAuthenticatedOrReadOnly]
#
#     def get(self, request):
#         return Response(data=ready_response, status=status.HTTP_200_OK)
#
#
# class HeavySpeedTestVS(APIView):
#     permission_classes = [IsAuthenticatedOrReadOnly]
#
#     def get(self, request):
#         return Response(data=heavy_ready_response, status=status.HTTP_200_OK)
#

# class HeavySpeedTestVS(viewsets.GenericViewSet, mixins.ListModelMixin, ):
#     permission_classes = [IsAuthenticatedOrReadOnly]
#     serializer_class = BookSerializer
#     queryset = Book.objects.filter(pk=1)

# def list(self, request, *args, **kwargs):
#     subjects =  Book.objects.filter(pk=1)
#     serializer = NarrationSubjectListSerializer({'subjects': subjects})
#     return Response(serializer.data)

# class HeavySpeedTestVS(mixins.ListModelMixin, mixins.CreateModelMixin, mixins.RetrieveModelMixin, mixins.UpdateModelMixin,
#              viewsets.GenericViewSet):
#     permission_classes = [PublicContentPermission]
#
#     def list(self, request, *args, **kwargs):
#         return Response(heavy_ready_response)
