import docx
from docx.shared import RGBColor
doc = docx.Document('kalamenoor/ahadis/test.docx')
copy_doc = docx.Document()
for para in doc.paragraphs:
    new_paragraph = copy_doc.add_paragraph()
    for run in para.runs:
        new_run = new_paragraph.add_run(run.text + ' ')
        new_run.font.color.rgb = run.font.color.rgb
copy_doc.save_narration('copy_doc111.docx')

print(doc.paragraphs[1].runs[1])

print(run.text)


print(run.font.color.rgb)

